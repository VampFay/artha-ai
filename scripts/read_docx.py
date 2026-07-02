import sys
from docx import Document

path = "/home/z/my-project/upload/FinSight AI Implementation Plan.docx"
doc = Document(path)

print("=" * 80)
print("DOCUMENT TEXT CONTENT")
print("=" * 80)

for i, para in enumerate(doc.paragraphs):
    if para.text.strip():
        style = para.style.name if para.style else "Normal"
        print(f"[{style}] {para.text}")

print("\n" + "=" * 80)
print("TABLES")
print("=" * 80)
for ti, table in enumerate(doc.tables):
    print(f"\n--- Table {ti+1} ---")
    for row in table.rows:
        cells = [cell.text.strip() for cell in row.cells]
        print(" | ".join(cells))
